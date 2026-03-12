"""
SWOT Generator + Investment Report Builder

Uses GenAI (Gemini) to produce:
  1. Comprehensive SWOT analysis (Strengths / Weaknesses / Opportunities / Threats)
  2. Downloadable Investment Report (.docx) with 8 sections

Falls back to deterministic generation if LLM is unavailable.
"""
import os
import json
from datetime import datetime
from typing import Dict, Any, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT


class SWOTGenerator:
    """Generate SWOT from scoring + research + triangulation data."""

    def generate_swot(
        self,
        company_name: str,
        sector: str,
        scoring: Dict[str, Any],
        financials: Dict[str, Any],
        research_bundle: Dict[str, Any],
        triangulation: Dict[str, Any],
        llm=None,
    ) -> Dict[str, Any]:
        """
        Produce SWOT dict.  If an LLM service is passed, use GenAI enrichment.
        """
        if llm:
            try:
                return self._llm_swot(
                    company_name, sector, scoring, financials,
                    research_bundle, triangulation, llm,
                )
            except Exception as e:
                print(f"[SWOT] LLM failed ({e}), falling back to deterministic")

        return self._deterministic_swot(
            company_name, sector, scoring, financials,
            research_bundle, triangulation,
        )

    # ------------------------------------------------------------------
    # LLM-based SWOT
    # ------------------------------------------------------------------
    def _llm_swot(self, company, sector, scoring, fin, research, tri, llm) -> Dict[str, Any]:
        credit_score = scoring.get("final_credit_score", 0)
        decision = scoring.get("decision", "")
        subs = scoring.get("sub_scores", {})
        tri_findings = tri.get("findings", [])
        signals = research.get("overall_signals", {}).get("signals", [])

        prompt = f"""Generate a comprehensive SWOT analysis in JSON for a credit assessment.

Company: {company}
Sector: {sector}
Credit Score: {credit_score}/100
Decision: {decision}

Five Cs Sub-Scores:
{json.dumps(subs, indent=2)}

Key Financial Metrics:
- DSCR: {fin.get('dscr', 'N/A')}
- Debt-to-Equity: {fin.get('debt_to_equity', 'N/A')}
- Revenue: ₹{fin.get('revenue_cr', 'N/A')} Cr
- Operating Cash Flow: ₹{fin.get('operating_cash_flow_cr', 'N/A')} Cr

Research Signals (top 6):
{json.dumps(signals[:6], indent=2)}

Triangulation Findings:
{json.dumps([{{"metric": f["metric"], "status": f["status"], "detail": f["detail"]}} for f in tri_findings[:8]], indent=2)}

Return JSON with this EXACT structure (3-5 items per category, each a single sentence):
{{
  "strengths": ["...", "..."],
  "weaknesses": ["...", "..."],
  "opportunities": ["...", "..."],
  "threats": ["...", "..."],
  "executive_summary": "A 2-3 sentence executive summary of the SWOT analysis."
}}

Use formal banking language. Be specific and data-driven. Return ONLY valid JSON."""

        raw = llm.generate_text(prompt, max_tokens=2000, temperature=0.3, json_mode=True)
        result = json.loads(raw)
        # Validate keys
        for key in ("strengths", "weaknesses", "opportunities", "threats"):
            if key not in result or not isinstance(result[key], list):
                raise ValueError(f"Missing or invalid key: {key}")
        return result

    # ------------------------------------------------------------------
    # Deterministic fallback
    # ------------------------------------------------------------------
    def _deterministic_swot(self, company, sector, scoring, fin, research, tri) -> Dict[str, Any]:
        strengths, weaknesses, opportunities, threats = [], [], [], []
        credit_score = scoring.get("final_credit_score", 0)
        subs = scoring.get("sub_scores", {})

        # ── Strengths ──
        if credit_score >= 70:
            strengths.append(f"Strong overall credit score of {credit_score}/100 indicates solid financial health")
        for dim, info in subs.items():
            if info.get("score", 0) >= 75:
                strengths.append(f"High {dim} score ({info['score']}/100) — above-average performance")

        dscr = fin.get("dscr", 0)
        if dscr >= 1.5:
            strengths.append(f"Strong debt service coverage ratio of {dscr:.2f}")
        de = fin.get("debt_to_equity", 999)
        if de <= 1.0:
            strengths.append(f"Conservative leverage with D/E ratio of {de:.2f}")

        ocf = fin.get("operating_cash_flow_cr", 0)
        if ocf > 0:
            strengths.append(f"Positive operating cash flow of ₹{ocf:.1f} Cr")

        mgmt = research.get("management", {}).get("overall_sentiment")
        if mgmt == "POSITIVE":
            strengths.append("Positive market reputation of promoters and management")

        if not strengths:
            strengths.append("Application submitted for review — standard assessment in progress")

        # ── Weaknesses ──
        if credit_score < 60:
            weaknesses.append(f"Credit score of {credit_score}/100 is below the approval threshold")
        for dim, info in subs.items():
            if info.get("score", 0) < 50:
                weaknesses.append(f"Low {dim} score ({info['score']}/100) needs attention")

        if dscr < 1.0:
            weaknesses.append(f"DSCR of {dscr:.2f} indicates insufficient cash flow for debt service")
        if de > 2.0:
            weaknesses.append(f"High leverage with D/E ratio of {de:.2f}")

        gst_var = fin.get("gst_vs_bank_variance", 0)
        if gst_var > 10:
            weaknesses.append(f"GST-Bank revenue variance of {gst_var:.1f}% raises data integrity concerns")

        legal = research.get("legal", {})
        if legal.get("case_count", 0) > 0:
            weaknesses.append(f"{legal['case_count']} legal case(s) found — litigation exposure")

        if not weaknesses:
            weaknesses.append("No significant weaknesses identified in current assessment")

        # ── Opportunities ──
        macro = research.get("macro_trends", {})
        if macro.get("overall_sentiment") == "POSITIVE":
            opportunities.append(f"Favorable macro conditions for the {sector} sector")
        for theme in macro.get("themes", [])[:2]:
            opportunities.append(f"Macro theme: {theme} could provide tailwinds")

        mkt = research.get("market_sentiment", {})
        if mkt.get("overall_sentiment") == "POSITIVE":
            opportunities.append("Positive market sentiment supports business growth outlook")

        opportunities.append(f"Continued investment in {sector} sector driven by policy support")
        if not opportunities:
            opportunities.append("Standard sector opportunity assessment in progress")

        # ── Threats ──
        if macro.get("overall_sentiment") == "NEGATIVE":
            threats.append("Challenging macro environment poses headwinds")

        news = research.get("news", {})
        for ra in news.get("risk_alerts", [])[:2]:
            threats.append(f"News risk: {ra}")

        reg = research.get("regulatory", {})
        if reg.get("has_issues"):
            threats.append("Regulatory compliance concerns identified")

        if mkt.get("overall_sentiment") == "NEGATIVE":
            threats.append("Negative market sentiment may affect valuation and access to capital")

        threats.append("Global economic uncertainty and interest rate movements")
        if not threats:
            threats.append("Standard macro and sector risks apply")

        summary = (
            f"{company} in the {sector} sector scores {credit_score}/100. "
            f"The analysis identifies {len(strengths)} strength(s) and {len(weaknesses)} weakness(es), "
            f"with {len(opportunities)} opportunities and {len(threats)} threats in the operating environment."
        )

        return {
            "strengths": strengths[:5],
            "weaknesses": weaknesses[:5],
            "opportunities": opportunities[:5],
            "threats": threats[:5],
            "executive_summary": summary,
        }


class InvestmentReportGenerator:
    """Generate a downloadable .docx investment report."""

    def __init__(self, output_dir: str = "./downloads/investment_reports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate_report(
        self,
        application_id: str,
        company_name: str,
        sector: str,
        swot: Dict[str, Any],
        recommendation: Dict[str, Any],
        triangulation: Dict[str, Any],
        scoring: Dict[str, Any],
        financials: Dict[str, Any],
        research_bundle: Dict[str, Any],
    ) -> str:
        """Generate .docx and return filepath."""
        doc = Document()
        self._set_styles(doc)

        self._cover(doc, company_name, sector, application_id)
        self._sec_executive_summary(doc, swot, recommendation)
        self._sec_swot(doc, swot)
        self._sec_financial_snapshot(doc, financials, scoring)
        self._sec_five_cs(doc, scoring)
        self._sec_research_signals(doc, research_bundle)
        self._sec_triangulation(doc, triangulation)
        self._sec_recommendation(doc, recommendation)
        self._sec_reasoning_chain(doc, recommendation)
        self._footer(doc)

        filename = f"{application_id}_investment_report.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    # ── Styles ────────────────────────────────────────────────────
    @staticmethod
    def _set_styles(doc):
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

    @staticmethod
    def _heading(doc, text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x16, 0x0B, 0x39)
        return h

    @staticmethod
    def _kv_table(doc, pairs):
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for k, v in pairs:
            row = table.add_row()
            row.cells[0].text = str(k)
            row.cells[1].text = str(v)
        doc.add_paragraph()

    @staticmethod
    def _bullet(doc, text):
        doc.add_paragraph(text, style="List Bullet")

    # ── Sections ──────────────────────────────────────────────────
    def _cover(self, doc, company, sector, app_id):
        title = doc.add_heading("INVESTMENT ANALYSIS REPORT", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run(f"\n{company}\n").bold = True
        p.add_run(f"Sector: {sector}\n").font.size = Pt(11)
        p.add_run(f"Application ID: {app_id}\n").font.size = Pt(9)
        p.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}").font.size = Pt(9)
        doc.add_page_break()

    def _sec_executive_summary(self, doc, swot, rec):
        self._heading(doc, "1. Executive Summary")
        doc.add_paragraph(swot.get("executive_summary", ""))
        doc.add_paragraph(rec.get("narrative", ""))

    def _sec_swot(self, doc, swot):
        self._heading(doc, "2. SWOT Analysis")
        for label, key in [("Strengths", "strengths"), ("Weaknesses", "weaknesses"),
                          ("Opportunities", "opportunities"), ("Threats", "threats")]:
            doc.add_heading(label, level=2)
            for item in swot.get(key, []):
                self._bullet(doc, item)

    def _sec_financial_snapshot(self, doc, fin, scoring):
        self._heading(doc, "3. Financial Snapshot")
        self._kv_table(doc, [
            ("Revenue", f"₹{fin.get('revenue_cr', 'N/A')} Cr"),
            ("DSCR", f"{fin.get('dscr', 'N/A')}"),
            ("Debt-to-Equity", f"{fin.get('debt_to_equity', 'N/A')}"),
            ("Current Ratio", f"{fin.get('current_ratio', 'N/A')}"),
            ("Operating Cash Flow", f"₹{fin.get('operating_cash_flow_cr', 'N/A')} Cr"),
            ("Net Worth", f"₹{fin.get('net_worth_cr', 'N/A')} Cr"),
            ("Credit Score", f"{scoring.get('final_credit_score', '-')}/100"),
            ("Risk Grade", scoring.get("risk_grade", "-")),
        ])

    def _sec_five_cs(self, doc, scoring):
        self._heading(doc, "4. Five Cs Evaluation")
        subs = scoring.get("sub_scores", {})
        expls = scoring.get("explanations", {})
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, t in enumerate(["Dimension", "Score", "Weight", "Assessment"]):
            hdr[i].text = t
        for dim in ("character", "capacity", "capital", "collateral", "conditions"):
            info = subs.get(dim, {})
            row = table.add_row().cells
            row[0].text = dim.capitalize()
            row[1].text = f"{info.get('score', '-')}/100"
            row[2].text = f"{int(info.get('weight', 0) * 100)}%"
            row[3].text = (expls.get(dim, ""))[:200]
        doc.add_paragraph()

    def _sec_research_signals(self, doc, research):
        self._heading(doc, "5. Secondary Research Signals")
        signals = research.get("overall_signals", {}).get("signals", [])
        if not signals:
            doc.add_paragraph("No external signals aggregated.")
            return
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, t in enumerate(["Source", "Signal", "Detail"]):
            hdr[i].text = t
        for s in signals:
            row = table.add_row().cells
            row[0].text = s.get("source", "")
            row[1].text = s.get("signal", "")
            row[2].text = (s.get("detail", ""))[:200]
        doc.add_paragraph()

    def _sec_triangulation(self, doc, tri):
        self._heading(doc, "6. Data Triangulation")
        summary = tri.get("summary", {})
        self._kv_table(doc, [
            ("Confirmed", summary.get("confirmed", 0)),
            ("Contradictions", summary.get("contradictions", 0)),
            ("Warnings", summary.get("warnings", 0)),
            ("Gaps", summary.get("gaps", 0)),
            ("Confidence Grade", summary.get("confidence_grade", "-")),
        ])
        for f in tri.get("findings", []):
            icon = {"CONFIRMED": "✓", "CONTRADICTION": "✗", "WARNING": "⚠", "GAP": "?"}.get(f["status"], "•")
            self._bullet(doc, f"{icon} {f['metric']}: {f['detail']}")

    def _sec_recommendation(self, doc, rec):
        self._heading(doc, "7. Recommendation")
        decision = rec.get("decision", "PENDING")
        dp = doc.add_paragraph()
        dp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        dr = dp.add_run(f"  {decision}  ")
        dr.bold = True
        dr.font.size = Pt(18)
        color_map = {"APPROVE": RGBColor(0, 128, 0), "REJECT": RGBColor(200, 0, 0)}
        dr.font.color.rgb = color_map.get(decision, RGBColor(200, 150, 0))

        loan = rec.get("loan_recommendation", {})
        rate = rec.get("interest_rate", {})
        self._kv_table(doc, [
            ("Credit Score", f"{rec.get('credit_score', '-')}/100"),
            ("Confidence", f"{rec.get('confidence_pct', '-')}%"),
            ("Recommended Limit", f"₹{loan.get('recommended_limit_cr', '-')} Cr"),
            ("Interest Rate", f"{rate.get('final_interest_rate', 'N/A')}%"
             if rate.get("offered") else "Not Offered"),
        ])

        # Gate results
        doc.add_heading("Decision Gates", level=2)
        for g in rec.get("gates", []):
            icon = "✓" if g["passed"] else "✗"
            self._bullet(doc, f"{icon} Gate {g['gate']} — {g['name']}: {g['verdict']} (Score {g['score']})")

    def _sec_reasoning_chain(self, doc, rec):
        self._heading(doc, "8. Reasoning Chain (Audit Trail)")
        doc.add_heading("Key Risks", level=2)
        for r in rec.get("key_risks", []):
            self._bullet(doc, f"[{r['source']}] {r['detail']}")
        doc.add_heading("Key Strengths", level=2)
        for r in rec.get("key_strengths", []):
            self._bullet(doc, f"[{r['source']}] {r['detail']}")

    def _footer(self, doc):
        doc.add_paragraph()
        fp = doc.add_paragraph()
        fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fr = fp.add_run("— End of Investment Analysis Report —")
        fr.font.size = Pt(9)
        fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        doc.add_paragraph()
        fp2 = doc.add_paragraph()
        fp2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fr2 = fp2.add_run(f"Generated by Intelli-Credit • {datetime.now().strftime('%d %B %Y %H:%M')}")
        fr2.font.size = Pt(8)
        fr2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
