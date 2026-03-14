"""
CAM (Credit Appraisal Memo) Generator

Produces a professional Word document (.docx) with 10 sections:
  1. Executive Summary
  2. Company Profile
  3. Industry Analysis
  4. Financial Analysis
  5. Bank Statement Analysis
  6. GST Compliance
  7. Litigation Check
  8. Five Cs Evaluation
  9. Risk Assessment
  10. Loan Recommendation

Optionally calls LLM (Gemini) for narrative enrichment.
"""
import os
from datetime import datetime
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT


class CAMGenerator:

    def __init__(self, output_dir: str = "./downloads/cam_reports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------
    def generate_cam(self, data: Dict[str, Any], llm_narrative: Optional[str] = None) -> str:
        """Generate a .docx CAM and return the file path."""
        doc = Document()
        self._set_styles(doc)

        self._sec1_executive_summary(doc, data, llm_narrative)
        self._sec2_company_profile(doc, data)
        self._sec3_industry_analysis(doc, data)
        self._sec4_financial_analysis(doc, data)
        self._sec5_bank_analysis(doc, data)
        self._sec6_gst_compliance(doc, data)
        self._sec7_litigation(doc, data)
        self._sec8_five_cs(doc, data)
        self._sec9_risk_assessment(doc, data)
        self._sec10_recommendation(doc, data)

        app_id = data.get("application_id", "UNKNOWN")
        filename = f"{app_id}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    # ------------------------------------------------------------------
    # Styling
    # ------------------------------------------------------------------
    @staticmethod
    def _set_styles(doc: Document):
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    @staticmethod
    def _add_heading(doc, text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        return h

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------
    @staticmethod
    def _kv_table(doc, pairs):
        """Add a 2-column key/value table."""
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"
        for key, val in pairs:
            row = table.add_row()
            row.cells[0].text = str(key)
            row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
            row.cells[1].text = str(val)
        doc.add_paragraph()  # spacer

    @staticmethod
    def _bullet(doc, text):
        doc.add_paragraph(text, style="List Bullet")

    # ------------------------------------------------------------------
    # Section 1 — Executive Summary
    # ------------------------------------------------------------------
    def _sec1_executive_summary(self, doc, data, llm_narrative=None):
        title = doc.add_heading("CREDIT APPRAISAL MEMORANDUM", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        self._add_heading(doc, "1. Executive Summary")

        scoring = data.get("scoring", {})
        company = data.get("company", {})
        loan = data.get("loan_recommendation", {})
        interest = data.get("interest_rate", {})

        # Use LLM narrative if available, else build deterministic
        if llm_narrative:
            doc.add_paragraph(llm_narrative)
        elif data.get("executive_summary"):
            doc.add_paragraph(data["executive_summary"])
        else:
            doc.add_paragraph(
                f"{company.get('company_name', 'N/A')} operating in the {company.get('sector', 'N/A')} sector "
                f"has been assessed for a credit facility of ₹{company.get('requested_limit_cr', 0)} Crores."
            )

        # Decision banner
        decision = scoring.get("decision", "PENDING")
        dp = doc.add_paragraph()
        dp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        dr = dp.add_run(f"  {decision}  ")
        dr.bold = True
        dr.font.size = Pt(16)
        color_map = {"APPROVE": RGBColor(0, 128, 0), "REJECT": RGBColor(200, 0, 0)}
        dr.font.color.rgb = color_map.get(decision, RGBColor(200, 150, 0))

        self._kv_table(doc, [
            ("Credit Score", f"{scoring.get('final_credit_score', '-')}/100"),
            ("Risk Grade", scoring.get("risk_grade", "-")),
            ("Recommended Limit", f"₹{loan.get('recommended_limit_cr', '-')} Cr"),
            ("Interest Rate", f"{interest.get('final_interest_rate', '-')}%"
             if interest.get("offered") else "Not Offered (Rejected)"),
        ])

    # ------------------------------------------------------------------
    # Section 2 — Company Profile
    # ------------------------------------------------------------------
    def _sec2_company_profile(self, doc, data):
        self._add_heading(doc, "2. Company Profile")
        c = data.get("company", {})
        self._kv_table(doc, [
            ("Company Name", c.get("company_name", "N/A")),
            ("CIN", c.get("mca_cin", "N/A")),
            ("Sector", c.get("sector", "N/A")),
            ("Requested Limit", f"₹{c.get('requested_limit_cr', 0)} Cr"),
            ("Application ID", data.get("application_id", "N/A")),
        ])

    # ------------------------------------------------------------------
    # Section 3 — Industry Analysis
    # ------------------------------------------------------------------
    def _sec3_industry_analysis(self, doc, data):
        self._add_heading(doc, "3. Industry Analysis")
        cond = data.get("conditions", {})
        doc.add_paragraph(
            cond.get("sector_narrative", "Sector analysis based on news, rating agencies, and regulatory environment.")
        )
        risk = cond.get("sector_risk_score", "-")
        doc.add_paragraph(f"Sector Risk Score: {risk}/100")

    # ------------------------------------------------------------------
    # Section 4 — Financial Analysis
    # ------------------------------------------------------------------
    def _sec4_financial_analysis(self, doc, data):
        self._add_heading(doc, "4. Financial Analysis")
        fin = data.get("financials", {})
        self._kv_table(doc, [
            ("Revenue (GST / AR)", f"₹{fin.get('revenue_cr', fin.get('gst_sales_cr', '-'))} Cr"),
            ("DSCR", f"{fin.get('dscr', '-'):.2f}" if isinstance(fin.get("dscr"), (int, float)) else "-"),
            ("Current Ratio", f"{fin.get('current_ratio', '-'):.2f}" if isinstance(fin.get("current_ratio"), (int, float)) else "-"),
            ("Debt-to-Equity", f"{fin.get('debt_to_equity', '-'):.2f}" if isinstance(fin.get("debt_to_equity"), (int, float)) else "-"),
            ("Net Worth", f"₹{fin.get('net_worth_cr', '-')} Cr" if fin.get("net_worth_cr") else "N/A"),
            ("Promoter Holding", f"{fin.get('promoter_holding_pct', 0):.1f}%" if fin.get("promoter_holding_pct") else "N/A"),
            ("Promoter Pledge", f"{fin.get('pledged_holding_pct', 0):.1f}%" if fin.get("pledged_holding_pct") or fin.get("promoter_holding_pct") else "N/A"),
            ("Top 10 Borrowings", f"{fin.get('top10_borrowings_pct', 0):.1f}% of borrowings" if fin.get("top10_borrowings_pct") else "N/A"),
            ("Short-term Liabilities", f"{fin.get('short_term_liabilities_pct_total_liabilities', 0):.1f}% of liabilities" if fin.get("short_term_liabilities_pct_total_liabilities") else "N/A"),
        ])

    # ------------------------------------------------------------------
    # Section 5 — Bank Statement Analysis
    # ------------------------------------------------------------------
    def _sec5_bank_analysis(self, doc, data):
        self._add_heading(doc, "5. Bank Statement Analysis")
        fin = data.get("financials", {})
        inflows  = fin.get("bank_inflows_cr", 0)
        outflows = fin.get("bank_outflows_cr", 0)
        self._kv_table(doc, [
            ("Total Inflows", f"₹{inflows} Cr"),
            ("Total Outflows", f"₹{outflows} Cr"),
            ("Net Cash Retention", f"₹{round(inflows - outflows, 2)} Cr"),
            ("Bounced Cheques", fin.get("bounced_cheques", 0)),
            ("Overdraft Instances", fin.get("overdraft_instances", 0)),
        ])

    # ------------------------------------------------------------------
    # Section 6 — GST Compliance
    # ------------------------------------------------------------------
    def _sec6_gst_compliance(self, doc, data):
        self._add_heading(doc, "6. GST Compliance")
        fin = data.get("financials", {})
        var = fin.get("gst_vs_bank_variance", 0)
        self._kv_table(doc, [
            ("GSTR-1 Sales", f"₹{fin.get('gst_sales_cr', '-')} Cr"),
            ("GSTR-3B Sales", f"₹{fin.get('gst_3b_sales_cr', fin.get('gst_sales_cr', '-'))} Cr"),
            ("Bank Inflows", f"₹{fin.get('bank_inflows_cr', '-')} Cr"),
            ("GST vs Bank Variance", f"{var:.1f}%"),
        ])
        if var > 10:
            doc.add_paragraph("⚠️ Significant variance — potential revenue inflation.").runs[0].font.color.rgb = RGBColor(200, 0, 0)
        elif var > 5:
            doc.add_paragraph("⚠️ Moderate variance — warrants monitoring.")

    # ------------------------------------------------------------------
    # Section 7 — Litigation Check
    # ------------------------------------------------------------------
    def _sec7_litigation(self, doc, data):
        self._add_heading(doc, "7. Litigation Check")
        research = data.get("research", {})
        lit = research.get("litigation_count", 0)
        doc.add_paragraph(f"Total cases found: {lit}")
        for case in research.get("litigation_cases", []):
            self._bullet(doc, case.get("summary", str(case)))

        if lit == 0:
            doc.add_paragraph("✅ No adverse litigation detected.")

    # ------------------------------------------------------------------
    # Section 8 — Five Cs Evaluation
    # ------------------------------------------------------------------
    def _sec8_five_cs(self, doc, data):
        self._add_heading(doc, "8. Five Cs Evaluation")
        scoring = data.get("scoring", {})
        subs = scoring.get("sub_scores", {})
        expls = scoring.get("explanations", {})

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, t in enumerate(["Dimension", "Score", "Weight", "Assessment"]):
            hdr[i].text = t
            hdr[i].paragraphs[0].runs[0].bold = True

        for c_name in ("character", "capacity", "capital", "collateral", "conditions"):
            info = subs.get(c_name, {})
            row = table.add_row().cells
            row[0].text = c_name.capitalize()
            row[1].text = f"{info.get('score', '-')}/100"
            row[2].text = f"{int(info.get('weight', 0) * 100)}%"
            row[3].text = expls.get(c_name, "")[:200]

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section 9 — Risk Assessment
    # ------------------------------------------------------------------
    def _sec9_risk_assessment(self, doc, data):
        self._add_heading(doc, "9. Risk Assessment")
        reasons = data.get("decision_reasons", [])
        fraud = data.get("fraud", {})
        fin = data.get("financials", {})

        if fraud:
            doc.add_paragraph(f"Fraud Verification Score: {fraud.get('combined_score', '-')}/100 "
                              f"(Risk: {fraud.get('overall_risk', '-')})")

        if fin.get("pledged_holding_pct", 0):
            doc.add_paragraph(
                f"Governance signal: promoter pledge at {fin.get('pledged_holding_pct', 0):.1f}%."
            )
        if fin.get("top10_borrowings_pct", 0):
            doc.add_paragraph(
                f"Liquidity signal: top 10 borrowings account for {fin.get('top10_borrowings_pct', 0):.1f}% of borrowings."
            )
        if fin.get("short_term_liabilities_pct_total_liabilities", 0):
            doc.add_paragraph(
                f"Liability structure: short-term liabilities are {fin.get('short_term_liabilities_pct_total_liabilities', 0):.1f}% of total liabilities."
            )

        if reasons:
            doc.add_heading("Key Risk Factors", level=2)
            for r in reasons:
                sign = "⊕" if r.get("impact") == "POSITIVE" else "⊖"
                self._bullet(doc, f"{sign} {r.get('text', str(r))}")
        else:
            doc.add_paragraph("Risk factors are detailed in the Five Cs assessment above.")

    # ------------------------------------------------------------------
    # Section 10 — Loan Recommendation
    # ------------------------------------------------------------------
    def _sec10_recommendation(self, doc, data):
        self._add_heading(doc, "10. Loan Recommendation")
        scoring  = data.get("scoring", {})
        loan     = data.get("loan_recommendation", {})
        interest = data.get("interest_rate", {})

        decision = scoring.get("decision", "PENDING")
        score    = scoring.get("final_credit_score", "-")

        dp = doc.add_paragraph()
        dp.add_run("Decision: ").bold = True
        dr = dp.add_run(decision)
        dr.bold = True

        self._kv_table(doc, [
            ("Credit Score", f"{score}/100"),
            ("Risk Grade", scoring.get("risk_grade", "-")),
            ("Recommended Loan Amount", f"₹{loan.get('recommended_limit_cr', '-')} Cr"),
            ("Revenue-Based Limit", f"₹{loan.get('revenue_limit_cr', '-')} Cr"),
            ("Cash-Flow-Based Limit", f"₹{loan.get('cashflow_limit_cr', '-')} Cr"),
            ("Collateral-Based Limit", f"₹{loan.get('collateral_limit_cr', '-')} Cr"),
            ("Interest Rate", f"{interest.get('final_interest_rate', '-')}%"
             if interest.get("offered") else "Not Offered"),
            ("Rate Category", interest.get("rate_category", "-")),
        ])

        # Explanation paragraph
        reasons = data.get("decision_reasons", [])
        if reasons:
            doc.add_heading("Explanation", level=2)
            neg = [r["text"] for r in reasons if r.get("impact") == "NEGATIVE"]
            pos = [r["text"] for r in reasons if r.get("impact") == "POSITIVE"]
            if pos:
                doc.add_paragraph("Strengths: " + "; ".join(pos[:3]))
            if neg:
                doc.add_paragraph("Concerns: " + "; ".join(neg[:3]))

        # Footer
        doc.add_paragraph()
        fp = doc.add_paragraph()
        fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fr = fp.add_run("— End of Credit Appraisal Memorandum —")
        fr.font.size = Pt(9)
        fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
